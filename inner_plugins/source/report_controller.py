from pathlib import Path
from typing import Union, List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import platform

from core._config import _const
from core.tooller.insert_ole_to_docx import insert_files_to_docx
from inner_plugins.source.report_table_row_register import get_row_info, BaseCell, EmptyCell, CellInfo

# 测试进度模版
TEST_PROGRESS = "%s\n上述所有模块，均已完成测试，测试进度100%%"
# 测试用例模版
TEST_CASE = "%s"
# BUG模版
TEST_BUG = "%s"
# 缺陷级别Mapping
BUG_SEVERITY_MAPPING = {
    "1": "一级问题",
    "2": "二级问题",
    "3": "三级问题",
    "4": "四级问题",
}
# 解决方案Mapping
BUG_RESOLUTION_MAPPING = {
    "": "未指定",
    "bydesign": "设计如此",
    "duplicate": "重复Bug",
    "external": "外部原因",
    "fixed": "已解决",
    "notrepro": "无法重现",
    "postponed": "延期处理",
    "willnotfix": "不予解决",
    "inside": "内部原因",
    "setup": "配置问题",
    "chargereq": "需求变更"
}


def collection_zen_tao_server(reference_data_object, data):
    # 项目名称
    reference_data_object["project_name"] = data.task["executionName"].split('/')[0].strip()
    # 测试周期
    if reference_data_object.get("time", None) is None:
        if reference_data_object["name"] == "生产环境":
            reference_data_object["time"] = datetime.now().strftime("%Y-%m-%d")
        else:
            reference_data_object["time"] = f"{data.task['begin']} - {data.task['end']}"
    # 测试版本
    reference_data_object["version"] = data.task["executionName"].split('/')[1].strip()
    # 标题
    reference_data_object["report_title"] = data.task["executionName"].replace('/', '-').replace(" ", "") + "-" + \
                                            reference_data_object[
                                                "name"] + "测试报告"
    # 缺陷
    bug_string = reference_data_object["version"] + f"共发现{data['bug']['total']}个问题，其中"
    for level, count in data['bug']["severity_mapping"].items():
        bug_string += BUG_SEVERITY_MAPPING[level] + str(count) + "个，"
    for resolution, count in data['bug']["resolution_mapping"].items():
        bug_string += BUG_RESOLUTION_MAPPING[resolution] + str(count) + "个，"
    bug_string = bug_string[:-1] + '。'
    reference_data_object["bugs"] = TEST_BUG % (bug_string,)
    reference_data_object["bug_file_path"] = data.bug_file_path
    return reference_data_object


def collection_xmind_plugin(reference_data_object, data):
    reference_data_object["test_cases"] = TEST_CASE % (data["case_result"].text,)
    return reference_data_object


def collection_excel_summary_plugin(reference_data_object, data):
    reference_data_object["tester"] = "，".join(list(data.pic_list))
    reference_data_object["model"] = "，".join([item["model"] for item in data.summary_data])
    reference_data_object["models"] = data.summary_data
    reference_data_object["test_progress"] = TEST_PROGRESS % (reference_data_object["model"],)
    reference_data_object["xmind_file_list"] = data.xmind_file_list

    return reference_data_object


def dispatch_info_collection(name: str, reference_data_object: dict, data):
    match name:
        case "ZenDaoServer":
            return collection_zen_tao_server(reference_data_object, data)
        case "XmindPlugin":
            return collection_xmind_plugin(reference_data_object, data)
        case "ExcelSummaryPlugin":
            return collection_excel_summary_plugin(reference_data_object, data)
        case _:
            return reference_data_object


class ReportController:

    def __init__(self, report_title: str = "", font_name: Union[str, None] = None, out_put_dir: str = None,
                 bug_file_path: str = None, xmind_file_list: List[str] = None, *args,
                 **kwargs):
        self.report_title = report_title
        self.font_name = font_name
        self.doc: Document = None
        self.table = None
        self.out_put_dir = out_put_dir
        self.bug_file_path = bug_file_path
        self.xmind_file_list = xmind_file_list

    def __enter__(self):
        self.doc = Document()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        insert_mapping = {
            _const.SYMBOL.DocxFileBugsInsertSymbol: [self.bug_file_path],
            _const.SYMBOL.DocxFileCasesInsertSymbol: self.xmind_file_list
        }
        self.doc = insert_files_to_docx(self.doc, insert_mapping)
        save_path = Path(self.out_put_dir) / (self.report_title + '.docx')
        self.doc.save(save_path)
        print(_const.NOTIFY.Success_Generation_File % (save_path,))

    def draw_title(self):
        secrets_title_paragraph = self.doc.add_paragraph()
        secrets_title = secrets_title_paragraph.add_run("机密▲3年")
        self.set_font_info(secrets_title, self.get_preferred_font(), 10, True)
        title_paragraph = self.doc.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title = title_paragraph.add_run(self.report_title)
        self.set_font_info(title, self.get_preferred_font(), 13, True)

    @staticmethod
    def set_font_info(paragraph, font_name, size, bold):
        paragraph.font.name = font_name
        paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        paragraph.font.size = Pt(size)
        paragraph.bold = bold

    def generate_report(self, row_list: List[List[Union[BaseCell, CellInfo]]]):
        self.draw_title()
        self.draw_table(len(row_list))
        for row_index, row in enumerate(row_list):
            for col_index, cell in enumerate(row):
                if isinstance(cell, EmptyCell):
                    continue
                if cell.is_title:
                    self.add_model_title(cell.value, row_index, col_index)
                else:
                    self.add_value(cell.value, row_index, col_index, cell.other_paragraphs)
                if cell.row_merge_count > 1:
                    self.merge_row(col_index, row_index, row_index + cell.row_merge_count - 1)
                if cell.col_merge_count > 1:
                    self.merge_col(row_index, col_index, col_index + cell.col_merge_count - 1)

    def draw_table(self, table_row_number):
        self.table = self.doc.add_table(rows=table_row_number, cols=4)
        self.table.style = "Table Grid"

    def merge_col(self, row_index, start, end):
        cell = self.table.cell(row_index, start)
        cell.merge(self.table.cell(row_index, end))

    def merge_row(self, col_index, start, end):
        cell = self.table.cell(start, col_index)
        cell.merge(self.table.cell(end, col_index))

    def add_model_title(self, text, row, col):
        cell = self.table.cell(row, col)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        self.set_font_info(run, self.get_preferred_font(), 10, True)

    def add_value(self, text, row, col, other_paragraphs: List[str]):
        cell = self.table.cell(row, col)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        self.set_font_info(run, self.get_preferred_font(), 10, False)
        if other_paragraphs is not None:
            for other_paragraph in other_paragraphs:
                new_paragraph = cell.add_paragraph()
                new_run = new_paragraph.add_run(other_paragraph)
                self.set_font_info(new_run, self.get_preferred_font(), 10, False)

    def get_preferred_font(self):
        if self.font_name is not None:
            return self.font_name
        system = platform.system()

        if system == 'Windows':
            return 'Microsoft YaHei'
        elif system == 'Darwin':
            return 'PingFang SC'
        elif system == 'Linux':
            return 'DejaVu Sans'
        return 'Arial'


def generation_report(data):
    with ReportController(**data) as rc:
        rc.generate_report(get_row_info(data))
